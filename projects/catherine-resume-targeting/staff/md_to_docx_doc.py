#!/usr/bin/env python3
"""Convert a long structured markdown document to a styled .docx.

General-purpose companion to md_to_docx.py (which handles one-page resumes).
This one handles full reports: headings h1-h4, bold/italic/code inline,
bullet and numbered lists, GFM pipe tables, blockquotes, and horizontal rules.

Usage:  python3 md_to_docx_doc.py <input.md> <output.docx>
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

BODY_FONT = "Calibri"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)      # deep blue for headings
QUOTE_GREY = RGBColor(0x44, 0x44, 0x44)

INLINE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[(.+?)\]\((.+?)\)")


def add_runs(paragraph, text, *, base_italic=False, base_bold=False,
             color=None, size=None):
    """Render inline markdown (**bold**, *italic*, `code`, [link](url)) as runs."""
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            _run(paragraph, text[pos:m.start()], base_bold, base_italic,
                 color, size, False)
        if m.group(1) is not None:
            _run(paragraph, m.group(1), True, base_italic, color, size, False)
        elif m.group(2) is not None:
            _run(paragraph, m.group(2), base_bold, True, color, size, False)
        elif m.group(3) is not None:
            _run(paragraph, m.group(3), base_bold, base_italic, color, size, True)
        elif m.group(4) is not None:
            _run(paragraph, m.group(4), base_bold, base_italic, color, size, False)
        pos = m.end()
    if pos < len(text):
        _run(paragraph, text[pos:], base_bold, base_italic, color, size, False)


def _run(paragraph, text, bold, italic, color, size, mono):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Consolas" if mono else BODY_FONT
    r.font.size = Pt(size if size else (9.5 if mono else 11))
    if color is not None:
        r.font.color.rgb = color


def shade_cell(cell, hex_fill):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(sh)


def hrule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    bd = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B0B0B0")
    bd.append(bottom)
    pPr.append(bd)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)


def add_heading(doc, text, level):
    sizes = {1: 20, 2: 15, 3: 12.5, 4: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    add_runs(p, text, base_bold=True, color=ACCENT, size=sizes.get(level, 11))
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    bd = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "1F4E79")
    bd.append(left)
    pPr.append(bd)
    add_runs(p, text, base_italic=True, color=QUOTE_GREY)


def add_table(doc, rows):
    cols = len(rows[0])
    t = doc.add_table(rows=len(rows), cols=cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = t.cell(ri, ci)
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            txt = row[ci] if ci < len(row) else ""
            add_runs(para, txt, base_bold=(ri == 0),
                     color=(RGBColor(0xFF, 0xFF, 0xFF) if ri == 0 else None),
                     size=10)
            if ri == 0:
                shade_cell(cell, "1F4E79")
            elif ri % 2 == 0:
                shade_cell(cell, "EEF2F7")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def split_table_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_sep_row(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def add_list_item(doc, text, ordered, level):
    style = ("List Number" if ordered else "List Bullet")
    if level >= 2:
        style += f" {min(level, 3)}"
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph(style=("List Number" if ordered else "List Bullet"))
    p.paragraph_format.space_after = Pt(2)
    add_runs(p, text)


def convert(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # horizontal rule
        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            hrule(doc)
            i += 1
            continue

        # heading
        hm = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if hm:
            add_heading(doc, hm.group(2).strip(), len(hm.group(1)))
            i += 1
            continue

        # table
        if stripped.startswith("|") and i + 1 < n and is_sep_row(lines[i + 1]):
            header = split_table_row(stripped)
            rows = [header]
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        # blockquote
        if stripped.startswith(">"):
            qtext = re.sub(r"^>\s?", "", stripped)
            i += 1
            while i < n and lines[i].strip().startswith(">"):
                qtext += " " + re.sub(r"^>\s?", "", lines[i].strip())
                i += 1
            add_blockquote(doc, qtext)
            continue

        # ordered list item
        om = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if om:
            level = len(om.group(1)) // 3 + 1
            add_list_item(doc, om.group(3), True, level)
            i += 1
            continue

        # bullet list item
        bm = re.match(r"^(\s*)[-*+]\s+(.*)$", line)
        if bm:
            level = len(bm.group(1)) // 2 + 1
            add_list_item(doc, bm.group(2), False, level)
            i += 1
            continue

        # plain paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        add_runs(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f"wrote {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        sys.exit("usage: md_to_docx_doc.py <input.md> <output.docx>")
    convert(sys.argv[1], sys.argv[2])
