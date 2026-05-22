#!/usr/bin/env python3
"""Convert a structured resume markdown file to a .docx matching the
original Catherine resume house style.

Markdown contract (see staff/RESUME_MD_FORMAT.md):
  TAGLINE: <one-line tagline>
  CONTACT: <contact line>
  # SECTION HEADER
  ## Role Title — Org | Location  •  Dates
  - bullet line
  plain line  -> body paragraph
Blank lines separate blocks.
"""
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAME = "Catherine Olver, EdD"


def make_run(p, text, *, bold=False, size=10.5):
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    return r


def build(md_path, docx_path):
    with open(md_path) as f:
        lines = f.read().splitlines()

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.75)
    sec.top_margin = sec.bottom_margin = Inches(0.75)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    tagline = ""
    contact = ""
    body = []
    for ln in lines:
        if ln.startswith("TAGLINE:"):
            tagline = ln[len("TAGLINE:"):].strip()
        elif ln.startswith("CONTACT:"):
            contact = ln[len("CONTACT:"):].strip()
        else:
            body.append(ln)

    # --- header block ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    make_run(p, NAME, bold=True, size=20)

    if tagline:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        make_run(p, tagline, size=11)

    if contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        make_run(p, contact, size=10)

    def add_section_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(5)
        make_run(p, text.upper(), bold=True, size=11)

    def add_role(text):
        # Title — Org | Location • Dates
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        rest = text
        # bold up to first ' — '
        if " — " in rest:
            title, rest = rest.split(" — ", 1)
        else:
            title, rest = rest, ""
        make_run(p, title, bold=True, size=11)
        if rest:
            # split org from " | location • dates"
            if " | " in rest:
                org, tail = rest.split(" | ", 1)
                make_run(p, " — " + org, size=11)
                make_run(p, "  |  " + tail, size=10)
            else:
                make_run(p, " — " + rest, size=11)

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        make_run(p, text, size=10.5)

    def add_body(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        make_run(p, text, size=10.5)

    for ln in body:
        s = ln.rstrip()
        if not s.strip():
            continue
        if s.startswith("## "):
            add_role(s[3:].strip())
        elif s.startswith("# "):
            add_section_header(s[2:].strip())
        elif s.lstrip().startswith("- "):
            add_bullet(s.lstrip()[2:].strip())
        else:
            add_body(s.strip())

    doc.save(docx_path)
    print(f"wrote {docx_path}")


if __name__ == "__main__":
    build(sys.argv[1], sys.argv[2])
